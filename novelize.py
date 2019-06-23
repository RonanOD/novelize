from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
import os, re, sys
###
# Script to add a front image to a novelize document. Front image file must be
# images/frontImage.jpg.
# Also add any images with "__*imageName*__" markup in a paragraph in novelize.
# All images go in the images folder
# Call script with doc name as argument: python novelize.py my-manuscript.docx
# Output is final-my-manuscript.docx with transformations included.
#
# Useful links:
# python-docx - https://python-docx.readthedocs.io/en/latest/index.html
# replacing text - https://www.quora.com/How-can-I-find-and-replace-text-in-a-Word-document-using-Python
###

# Variables and Constants
docName = ''
finalDocName = ''
frontImage = 'images/frontImage.jpg'
imgRegex = re.compile(r"__\*(.*)\*__")
fontRegex = re.compile(r"_f_(.*)_fS(\d+)_")
breakRegex = re.compile(r"_p_PAGE_BREAK_p_")
removeRegex = re.compile(r"__REMOVE__")
emdashRegex = re.compile(r"(-- )")
imageWidth = 6.5

def main():
    # Start up
    if len(sys.argv) < 2:
        print("Run script with document name as argument: novelize.py my-manuscript.docx")
        exit()
    else:
        docName = sys.argv[1]
        finalDocName = 'final-' + docName

    # Remove final doc if necessary
    if os.path.exists(finalDocName):
        os.remove(finalDocName)
    # Open the document
    document = Document(docName)
    # Add front image
    para = document.paragraphs[0].insert_paragraph_before()
    run = para.add_run()
    run.add_picture(frontImage, width=Inches(imageWidth))
    run.add_break(WD_BREAK.PAGE)

    # Look for __*imageName*__ in a paragraph on its own 
    # and replace with imageName in same folder:
    for p in document.paragraphs:
        imgMatch = imgRegex.search(p.text)
        if imgMatch:
            imageFile = "images/" + imgMatch.group(1)
            for run in p.runs: # delete the text
                run.text = ""
            lastRun = p.add_run()
            lastRun.add_picture(imageFile, width=Inches(imageWidth))
            continue
        
        fontMatch = fontRegex.search(p.text)
        if fontMatch:
            # Change the font
            txt = fontMatch.group(1).strip("_")
            number = int(fontMatch.group(2))
            for run in p.runs: # delete the text
                run.text = ""
            lastRun = p.add_run()
            lastRun.font.size = Pt(number)
            lastRun.text = txt
            continue

        pbMatch = breakRegex.match(p.text)
        if pbMatch:
            # Insert a page break
            for run in p.runs: # delete the text
                run.text = ""
            lastRun = p.add_run()
            lastRun.add_break(WD_BREAK.PAGE)
            continue

        emdashMatch = emdashRegex.search(p.text)
        if emdashMatch:
            # Replace -- with an emdash
            p.text = p.text.replace("-- ", "– ")
            print(p.text)
            continue

        removeMatch = removeRegex.search(p.text)
        if removeMatch:
            removeText(p)
            continue

    # Finished. Save the final document
    document.save(finalDocName)

def removeText(p):
    for run in p.runs: # delete the text
        run.text = ""
    lastRun = p.add_run
    return lastRun

if __name__== "__main__":
    main()